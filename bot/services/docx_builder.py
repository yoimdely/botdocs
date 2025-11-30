from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Dict

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .legal import DISCLAIMER_TEXT
from .templates_loader import TemplateLoader


class DocxBuilder:
    def __init__(self, template_loader: TemplateLoader) -> None:
        self.template_loader = template_loader

    def build(self, template_name: str, context: Dict[str, str]) -> BytesIO:
        document = Document()

        section = document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(14)
        normal_style.paragraph_format.first_line_indent = Cm(1.25)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        title_style = document.styles.add_style("GOSTTitle", WD_STYLE_TYPE.PARAGRAPH, builtin=False)
        title_style.font.name = "Times New Roman"
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

        meta_style = document.styles.add_style("GOSTMeta", WD_STYLE_TYPE.PARAGRAPH, builtin=False)
        meta_style.font.name = "Times New Roman"
        meta_style.font.size = Pt(12)
        meta_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_style.paragraph_format.space_after = Pt(6)
        meta_style.paragraph_format.first_line_indent = Cm(0)

        disclaimer_style = document.styles.add_style(
            "GOSTDisclaimer", WD_STYLE_TYPE.PARAGRAPH, builtin=False
        )
        disclaimer_style.font.name = "Times New Roman"
        disclaimer_style.font.size = Pt(10)
        disclaimer_style.paragraph_format.first_line_indent = Cm(0)
        disclaimer_style.paragraph_format.line_spacing = 1.0
        disclaimer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        footer_style = document.styles.add_style(
            "GOSTFooter", WD_STYLE_TYPE.PARAGRAPH, builtin=False
        )
        footer_style.font.name = "Times New Roman"
        footer_style.font.size = Pt(12)
        footer_style.paragraph_format.first_line_indent = Cm(0)
        footer_style.paragraph_format.line_spacing = 1.15
        footer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        rendered = self.template_loader.render(template_name, context)

        lines = [line.rstrip() for line in rendered.split("\n")]

        first_content_added = False
        for line in lines:
            if not line.strip():
                continue

            if not first_content_added:
                document.add_paragraph(line.strip(), style=title_style)
                first_content_added = True
                continue

            if line.lower().startswith("г. "):
                document.add_paragraph(line.strip(), style=meta_style)
                continue

            document.add_paragraph(line.strip())

        document.add_paragraph()
        document.add_paragraph(
            f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            style=footer_style,
        )
        document.add_paragraph(
            "Подпись стороны: _____________________", style=footer_style
        )
        document.add_paragraph()
        disclaimer = document.add_paragraph(DISCLAIMER_TEXT, style=disclaimer_style)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer
